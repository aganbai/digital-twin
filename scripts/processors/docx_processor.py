"""DOCX文件处理器"""

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def process_docx(filepath):
    """解析DOCX文件，提取文本内容"""
    if not HAS_DOCX:
        raise ImportError("python-docx 未安装，请运行: pip install python-docx")

    doc = Document(filepath)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))

    content = "\n".join(text_parts)
    if not content.strip():
        raise ValueError("DOCX文件中未能提取到文本内容")
    return content
